from __future__ import annotations

import html
import json
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, Request, Response, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.config import PROJECT_ROOT, get_settings
from app.factory import build_runtime
from app.llm import TextGuard
from app.services.editorial_department import EditorialDepartment
from app.storage.repository import POV_LABELS, json_safe


EXPORT_ROOT = PROJECT_ROOT / "exports"
EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
UPLOAD_ROOT = EXPORT_ROOT

app = FastAPI(title="fiction_architect")
app.mount("/static", StaticFiles(directory=str(PROJECT_ROOT / "app" / "static")), name="static")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_ROOT)), name="uploads")
templates = Jinja2Templates(directory=str(PROJECT_ROOT / "app" / "templates"))


@app.middleware("http")
async def no_store_html(request: Request, call_next):
    response: Response = await call_next(request)
    content_type = response.headers.get("content-type", "")
    if "text/html" in content_type:
        response.headers["Cache-Control"] = "no-store, max-age=0"
    return response


def runtime():
    return build_runtime()


def render_error(request: Request, message: str, status_code: int = 400):
    return templates.TemplateResponse(request, "error.html", {"request": request, "message": message}, status_code=status_code)


def _menu(book_id: int) -> list[dict[str, str]]:
    return [
        {"label": "开书设定", "href": f"/books/{book_id}/setup", "hint": "全书结构化大纲、预估字数和基础设定"},
        {"label": "卷纲", "href": f"/books/{book_id}/outline", "hint": "卷纲、单元纲、章节细纲入口"},
        {"label": "作者设置", "href": f"/books/{book_id}/author", "hint": "当前书专属文风和写作规则"},
        {"label": "编辑设置", "href": f"/books/{book_id}/editor", "hint": "当前书专属平台审稿规则"},
        {"label": "连续性工作室", "href": f"/books/{book_id}/continuity", "hint": "记忆、事实池、漂移审计"},
        {"label": "新建章节", "href": f"/books/{book_id}/chapter-batches/new", "hint": "按当前单元推荐章节数生成"},
        {"label": "正文/导出", "href": f"/books/{book_id}/chapters", "hint": "确认正文并导出 DOCX"},
    ]


def _book_context(repo, book_id: int, request: Request, **extra: Any) -> dict[str, Any]:
    book = repo.get_book(book_id)
    if book is None:
        return {"request": request, "missing": True}
    return {
        "request": request,
        "book": book,
        "record": repo.get_book_record(book_id) or {},
        "menu": _menu(book_id),
        "pov_labels": POV_LABELS,
        **extra,
    }


@app.get("/", response_class=HTMLResponse)
def dashboard(request: Request):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "dashboard.html", {"request": request, **repo.dashboard()})


@app.get("/books", response_class=HTMLResponse)
def books(request: Request):
    return dashboard(request)


@app.get("/books/create", response_class=HTMLResponse)
def create_book_page(request: Request):
    return templates.TemplateResponse(request, "book_create.html", {"request": request, "pov_labels": POV_LABELS})


@app.post("/books/create")
def create_book(
    request: Request,
    title: str = Form(...),
    premise: str = Form(""),
    genre: str = Form(""),
    market_channel: str = Form("番茄/男频"),
    target_reader: str = Form(""),
    pov_policy: str = Form("third_limited"),
    target_chars_min: int = Form(2200),
    target_chars_max: int = Form(3200),
    estimated_total_words: int = Form(1000000),
    story_mainline: str = Form(""),
    worldbuilding: str = Form(""),
    book_outline: str = Form(""),
    imported_outline: str = Form(""),
    characters_text: str = Form(""),
):
    repo, _ = runtime()
    try:
        book_id = repo.create_book(
            title=title,
            premise=premise,
            genre=genre,
            market_channel=market_channel,
            target_reader=target_reader,
            pov_policy=pov_policy,
            target_chars_min=target_chars_min,
            target_chars_max=target_chars_max,
            estimated_total_words=estimated_total_words,
            story_mainline=story_mainline,
            worldbuilding=worldbuilding,
            book_outline=book_outline,
            imported_outline=imported_outline,
            characters_text=characters_text,
        )
    except ValueError as exc:
        return render_error(request, str(exc))
    return RedirectResponse(f"/books/{book_id}/setup?created=1", status_code=303)


@app.get("/archive", response_class=HTMLResponse)
def archive(request: Request):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "archive.html", {"request": request, "books": repo.list_books("archived")})


@app.post("/books/{book_id}/archive")
def archive_book(book_id: int):
    repo, _ = runtime()
    repo.archive_book(book_id)
    return RedirectResponse("/", status_code=303)


@app.post("/books/{book_id}/restore")
def restore_book(request: Request, book_id: int):
    repo, _ = runtime()
    try:
        repo.restore_book(book_id)
    except ValueError as exc:
        return render_error(request, str(exc))
    return RedirectResponse(f"/books/{book_id}", status_code=303)


@app.post("/books/{book_id}/delete")
def delete_book(book_id: int, confirm: str = Form("")):
    repo, _ = runtime()
    if confirm == "DELETE":
        repo.delete_book_permanently(book_id)
    return RedirectResponse("/archive", status_code=303)


@app.get("/books/{book_id}", response_class=HTMLResponse)
def book_detail(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(
        repo,
        book_id,
        request,
        controls=repo.get_book_controls(book_id),
        volumes=repo.list_volumes(book_id),
        arcs=repo.list_arcs(book_id),
        plans=repo.list_chapter_plan_rows(book_id),
        batches=repo.list_chapter_batches(book_id),
        bodies=repo.list_chapter_bodies(book_id),
        exports=repo.list_exports(book_id),
        events=repo.list_events(book_id),
    )
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "book_detail.html", context)


@app.post("/books/{book_id}/cover")
def upload_cover(book_id: int, cover: UploadFile = File(...)):
    repo, _ = runtime()
    suffix = Path(cover.filename or "cover.bin").suffix.lower()
    if suffix not in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
        suffix = ".bin"
    safe_name = f"book_{book_id}{suffix}"
    target = UPLOAD_ROOT / safe_name
    target.write_bytes(cover.file.read())
    repo.update_cover(book_id, f"/uploads/{safe_name}")
    return RedirectResponse(f"/books/{book_id}", status_code=303)


@app.get("/books/{book_id}/setup", response_class=HTMLResponse)
def setup_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(repo, book_id, request, created=request.query_params.get("created") == "1")
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "book_setup.html", context)


@app.post("/books/{book_id}/setup")
def update_setup(
    request: Request,
    book_id: int,
    title: str = Form(...),
    premise: str = Form(""),
    genre: str = Form(""),
    market_channel: str = Form(""),
    target_reader: str = Form(""),
    pov_policy: str = Form("third_limited"),
    target_chars_min: int = Form(2200),
    target_chars_max: int = Form(3200),
    estimated_total_words: int = Form(1000000),
    story_mainline: str = Form(""),
    worldbuilding: str = Form(""),
    book_outline: str = Form(""),
    imported_outline: str = Form(""),
):
    repo, _ = runtime()
    try:
        repo.update_book_setup(
            book_id,
            {
                "title": title,
                "premise": premise,
                "genre": genre,
                "market_channel": market_channel,
                "target_reader": target_reader,
                "pov_policy": pov_policy,
                "target_chars_min": target_chars_min,
                "target_chars_max": target_chars_max,
                "estimated_total_words": estimated_total_words,
                "story_mainline": story_mainline,
                "worldbuilding": worldbuilding,
                "book_outline": book_outline,
                "imported_outline": imported_outline,
            },
        )
    except ValueError as exc:
        return render_error(request, str(exc))
    return RedirectResponse(f"/books/{book_id}", status_code=303)


@app.post("/books/{book_id}/setup/lock-outline")
def lock_outline(book_id: int):
    repo, _ = runtime()
    repo.lock_book_outline(book_id)
    return RedirectResponse(f"/books/{book_id}/setup", status_code=303)


@app.post("/books/{book_id}/setup/unlock-outline")
def unlock_outline(book_id: int):
    repo, _ = runtime()
    repo.unlock_book_outline(book_id)
    return RedirectResponse(f"/books/{book_id}/setup", status_code=303)


@app.get("/books/{book_id}/outline", response_class=HTMLResponse)
def outline_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(
        repo,
        book_id,
        request,
        controls=repo.get_book_controls(book_id),
        volumes=repo.list_volumes(book_id),
        arcs=repo.list_arcs(book_id),
        plans=repo.list_chapter_plan_rows(book_id),
    )
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "outline.html", context)


@app.post("/books/{book_id}/volumes/{volume_id}")
def update_volume(
    book_id: int,
    volume_id: int,
    title: str = Form(...),
    goal: str = Form(""),
    estimated_words: int = Form(200000),
    core_conflict: str = Form(""),
    stage_payoff: str = Form(""),
    character_progression: str = Form(""),
    foreshadowing_plan: str = Form(""),
    start_chapter: int = Form(1),
    end_chapter: int = Form(1),
):
    repo, _ = runtime()
    repo.update_volume(book_id, volume_id, locals())
    return RedirectResponse(f"/books/{book_id}/outline", status_code=303)


@app.post("/books/{book_id}/arcs/{arc_id}")
def update_arc(
    book_id: int,
    arc_id: int,
    title: str = Form(...),
    goal: str = Form(""),
    pressure: str = Form(""),
    cause: str = Form(""),
    process: str = Form(""),
    result: str = Form(""),
    payoff: str = Form(""),
    character_change: str = Form(""),
    foreshadowing_progress: str = Form(""),
    recommended_chapters: int = Form(5),
    start_chapter: int = Form(1),
    end_chapter: int = Form(1),
    status: str = Form("planned"),
):
    repo, _ = runtime()
    repo.update_arc(book_id, arc_id, locals())
    return RedirectResponse(f"/books/{book_id}/outline", status_code=303)


@app.get("/books/{book_id}/author", response_class=HTMLResponse)
def book_author_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(repo, book_id, request, controls=repo.get_book_controls(book_id), authors=repo.list_profiles("authors"))
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "book_author.html", context)


@app.post("/books/{book_id}/author")
def update_book_author(book_id: int, name: str = Form(""), genre: str = Form(""), pov_preference: str = Form("third_limited"), sentence_rhythm: str = Form(""), dialogue_style: str = Form(""), payoff_preference: str = Form(""), forbidden_items: str = Form(""), prompt_rules: str = Form(""), sample_summary: str = Form("")):
    repo, _ = runtime()
    repo.update_book_author_profile(book_id, locals())
    return RedirectResponse(f"/books/{book_id}", status_code=303)


@app.post("/books/{book_id}/author/from-resource")
def assign_author_resource(book_id: int, author_profile_id: int = Form(...)):
    repo, _ = runtime()
    controls = repo.get_book_controls(book_id)
    editor_id = int((controls["book"] or {}).get("editor_profile_id") or 0)
    repo.assign_resources(book_id, author_profile_id, editor_id)
    return RedirectResponse(f"/books/{book_id}/author", status_code=303)


@app.get("/books/{book_id}/editor", response_class=HTMLResponse)
def book_editor_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(repo, book_id, request, controls=repo.get_book_controls(book_id), editors=repo.list_profiles("editors"))
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "book_editor.html", context)


@app.post("/books/{book_id}/editor")
def update_book_editor(book_id: int, name: str = Form(""), platform: str = Form(""), word_count_rule: str = Form(""), pov_rule: str = Form(""), structure_rule: str = Form(""), payoff_rule: str = Form(""), pollution_rule: str = Form(""), reject_threshold: int = Form(1)):
    repo, _ = runtime()
    repo.update_book_editor_profile(book_id, locals())
    return RedirectResponse(f"/books/{book_id}", status_code=303)


@app.post("/books/{book_id}/editor/from-resource")
def assign_editor_resource(book_id: int, editor_profile_id: int = Form(...)):
    repo, _ = runtime()
    controls = repo.get_book_controls(book_id)
    author_id = int((controls["book"] or {}).get("author_profile_id") or 0)
    repo.assign_resources(book_id, author_id, editor_profile_id)
    return RedirectResponse(f"/books/{book_id}/editor", status_code=303)


@app.get("/books/{book_id}/chapter-batches/new", response_class=HTMLResponse)
def new_batch_page(request: Request, book_id: int):
    repo, _ = runtime()
    current_unit = repo.get_current_unit(book_id)
    context = _book_context(
        repo,
        book_id,
        request,
        volumes=repo.list_volumes(book_id),
        arcs=repo.list_arcs(book_id),
        current_unit=current_unit,
    )
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "chapter_batch_new.html", context)


@app.post("/books/{book_id}/chapter-batches/new")
def create_batch(request: Request, book_id: int, chapter_count: int = Form(0), arc_id: int = Form(0), volume_id: int = Form(0)):
    if chapter_count and (chapter_count < 1 or chapter_count > 20):
        return render_error(request, "一次最多只能新建 20 章，最少 1 章。")
    repo, _ = runtime()
    batch_id = repo.create_chapter_batch(book_id, chapter_count or None, volume_id or None, arc_id or None)
    return RedirectResponse(f"/books/{book_id}/chapter-batches/{batch_id}", status_code=303)


@app.get("/books/{book_id}/chapter-batches/{batch_id}", response_class=HTMLResponse)
def batch_page(request: Request, book_id: int, batch_id: int):
    repo, _ = runtime()
    batch = repo.get_chapter_batch(batch_id)
    if not batch or int(batch["book_id"]) != book_id:
        return render_error(request, "章节批次不存在。", 404)
    context = _book_context(
        repo,
        book_id,
        request,
        batch=batch,
        plans=repo.list_chapter_plan_rows(book_id, batch_id),
        bodies={row["chapter_no"]: row for row in repo.list_chapter_bodies(book_id)},
        artifacts_by_chapter={row["chapter_no"]: row for row in repo.list_artifacts(book_id) if row["artifact_type"] == "generation_error" and row["status"] == "failed"},
    )
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "chapter_batch.html", context)


@app.post("/books/{book_id}/chapter-batches/{batch_id}/generate")
def generate_batch(book_id: int, batch_id: int):
    _, pipe = runtime()
    try:
        pipe.generate_batch(book_id, batch_id)
    except Exception as exc:
        repo, _ = runtime()
        repo.update_batch_status(batch_id, "failed", "生成失败，可重试。", f"{type(exc).__name__}: {exc}")
    return RedirectResponse(f"/books/{book_id}/chapter-batches/{batch_id}", status_code=303)


@app.get("/books/{book_id}/chapters", response_class=HTMLResponse)
def chapters_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(
        repo,
        book_id,
        request,
        plans=repo.list_chapter_plan_rows(book_id),
        batches=repo.list_chapter_batches(book_id),
        bodies={row["chapter_no"]: row for row in repo.list_chapter_bodies(book_id)},
        exports=repo.list_exports(book_id),
    )
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "chapters.html", context)


@app.get("/books/{book_id}/chapters/{chapter_no}", response_class=HTMLResponse)
def chapter_detail(request: Request, book_id: int, chapter_no: int):
    repo, _ = runtime()
    plan = repo.get_chapter_plan_row(book_id, chapter_no)
    if plan is None:
        return render_error(request, "章节卡片不存在。", 404)
    context = _book_context(repo, book_id, request, plan=plan, body=repo.get_chapter_body(book_id, chapter_no), artifacts=repo.list_artifacts(book_id, chapter_no), task=repo.get_rewrite_task(book_id, chapter_no))
    return templates.TemplateResponse(request, "chapter_detail.html", context)


@app.post("/books/{book_id}/chapters/{chapter_no}/plan")
def update_chapter_plan(book_id: int, chapter_no: int, title: str = Form(...), objective: str = Form(...), allowed_reveals: str = Form(""), forbidden_reveals: str = Form(""), pace_limit: str = Form(""), plot_summary: str = Form(""), target_chars: int = Form(2600)):
    repo, _ = runtime()
    repo.update_chapter_plan(book_id, chapter_no, title, objective, allowed_reveals, forbidden_reveals, pace_limit, plot_summary, target_chars)
    plan = repo.get_chapter_plan_row(book_id, chapter_no) or {}
    batch_id = plan.get("batch_id")
    target = f"/books/{book_id}/chapter-batches/{batch_id}" if batch_id else f"/books/{book_id}/outline"
    return RedirectResponse(target, status_code=303)


@app.post("/books/{book_id}/chapters/{chapter_no}/generate")
def generate_chapter(book_id: int, chapter_no: int):
    _, pipe = runtime()
    pipe.generate_chapter(book_id, chapter_no)
    return RedirectResponse(f"/books/{book_id}/chapters/{chapter_no}", status_code=303)


@app.post("/books/{book_id}/chapters/{chapter_no}/rewrite")
def rewrite_chapter(book_id: int, chapter_no: int):
    _, pipe = runtime()
    pipe.rewrite_and_review(book_id, chapter_no)
    return RedirectResponse(f"/books/{book_id}/chapters/{chapter_no}", status_code=303)


@app.post("/books/{book_id}/chapters/{chapter_no}/delete")
def delete_chapter(request: Request, book_id: int, chapter_no: int, confirm: str = Form("")):
    repo, _ = runtime()
    if confirm != "DELETE":
        return render_error(request, "请输入 DELETE 确认删除章节。")
    result = repo.delete_chapter(book_id, chapter_no)
    batch_id = result.get("batch_id")
    target = f"/books/{book_id}/chapter-batches/{batch_id}" if batch_id else f"/books/{book_id}/chapters"
    return RedirectResponse(target, status_code=303)


@app.post("/books/{book_id}/chapters/{chapter_no}/confirm")
def confirm_chapter(request: Request, book_id: int, chapter_no: int, body: str = Form("")):
    repo, _ = runtime()
    try:
        repo.confirm_chapter_body(book_id, chapter_no, body)
    except ValueError as exc:
        return render_error(request, str(exc))
    return RedirectResponse(f"/books/{book_id}/chapters", status_code=303)


@app.get("/books/{book_id}/continuity", response_class=HTMLResponse)
def continuity_page(request: Request, book_id: int):
    repo, _ = runtime()
    context = _book_context(repo, book_id, request, memories=repo.list_memories(book_id), atoms=repo.list_atoms(book_id), logs=repo.list_retrieval_logs(book_id))
    if context.get("missing"):
        return render_error(request, "作品不存在。", 404)
    return templates.TemplateResponse(request, "continuity.html", context)


@app.get("/books/{book_id}/continuity/memories/{memory_id}/edit", response_class=HTMLResponse)
def memory_edit_page(request: Request, book_id: int, memory_id: int):
    repo, _ = runtime()
    memory = next((row for row in repo.list_memories(book_id, limit=1000) if int(row["id"]) == memory_id), None)
    if memory is None:
        return render_error(request, "记忆版本不存在。", 404)
    return templates.TemplateResponse(request, "memory_edit.html", _book_context(repo, book_id, request, memory=memory))


@app.post("/books/{book_id}/continuity/memories/{memory_id}/edit")
def update_memory(book_id: int, memory_id: int, content: str = Form(...)):
    repo, _ = runtime()
    repo.update_memory_content(memory_id, content)
    return RedirectResponse(f"/books/{book_id}/continuity", status_code=303)


@app.post("/books/{book_id}/continuity/compress-period")
def compress_period(book_id: int, memory_type: str = Form(...)):
    _, pipe = runtime()
    pipe.continuity.compress_period(book_id, memory_type)
    return RedirectResponse(f"/books/{book_id}/continuity", status_code=303)


@app.get("/books/{book_id}/continuity/atoms", response_class=HTMLResponse)
def atoms_page(request: Request, book_id: int):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "atoms.html", _book_context(repo, book_id, request, atoms=repo.list_atoms(book_id)))


@app.post("/books/{book_id}/continuity/atoms/{atom_id}/{status}")
def atom_status(book_id: int, atom_id: int, status: str):
    repo, _ = runtime()
    if status in {"approved", "rejected"}:
        repo.update_atom_status(atom_id, status)
    return RedirectResponse(f"/books/{book_id}/continuity/atoms", status_code=303)


@app.get("/books/{book_id}/continuity/drift", response_class=HTMLResponse)
def drift_page(request: Request, book_id: int):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "drift.html", _book_context(repo, book_id, request, reports=repo.list_drift_reports(book_id)))


@app.post("/books/{book_id}/continuity/drift")
def create_drift(book_id: int, chapter_no: int = Form(0)):
    _, pipe = runtime()
    pipe.continuity.drift_check(book_id, chapter_no or None)
    return RedirectResponse(f"/books/{book_id}/continuity/drift", status_code=303)


@app.get("/books/{book_id}/continuity/logs", response_class=HTMLResponse)
def retrieval_logs_page(request: Request, book_id: int):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "retrieval_logs.html", _book_context(repo, book_id, request, logs=repo.list_retrieval_logs(book_id)))


@app.get("/resources/{kind}", response_class=HTMLResponse)
def resources_page(request: Request, kind: str):
    repo, _ = runtime()
    if kind not in {"authors", "editors"}:
        return render_error(request, "资源类型不存在。", 404)
    return templates.TemplateResponse(request, "resources.html", {"request": request, "kind": kind, "profiles": repo.list_profiles(kind), "edit_profile": None, "json_text": ""})


@app.get("/resources/{kind}/{profile_id}", response_class=HTMLResponse)
def edit_resource_page(request: Request, kind: str, profile_id: int):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "resources.html", {"request": request, "kind": kind, "profiles": repo.list_profiles(kind), "edit_profile": json_safe(repo.get_profile(kind, profile_id)), "json_text": ""})


@app.post("/resources/{kind}/save")
def save_resource(request: Request, kind: str, profile_id: int = Form(0), payload: str = Form("")):
    repo, _ = runtime()
    try:
        saved_id = repo.save_profile(kind, json.loads(payload), profile_id or None)
    except (json.JSONDecodeError, TypeError, ValueError) as exc:
        return render_error(request, f"资源 JSON 无法保存：{exc}")
    return RedirectResponse(f"/resources/{kind}/{saved_id}", status_code=303)


@app.post("/resources/{kind}/delete")
def delete_resource(kind: str, profile_id: int = Form(...)):
    repo, _ = runtime()
    repo.delete_profile(kind, profile_id)
    return RedirectResponse(f"/resources/{kind}", status_code=303)


@app.post("/resources/{kind}/import")
def import_resource(request: Request, kind: str, json_text: str = Form(...)):
    repo, _ = runtime()
    try:
        repo.import_profiles(kind, json_text)
    except (json.JSONDecodeError, TypeError, ValueError) as exc:
        return render_error(request, f"资源 JSON 无法导入：{exc}")
    return RedirectResponse(f"/resources/{kind}", status_code=303)


@app.get("/resources/{kind}/export/json")
def export_resources(kind: str):
    repo, _ = runtime()
    path = EXPORT_ROOT / f"{kind}_resources.json"
    path.write_text(json.dumps(json_safe(repo.list_profiles(kind)), ensure_ascii=False, indent=2), encoding="utf-8")
    return FileResponse(path, filename=path.name, media_type="application/json")


@app.get("/debug", response_class=HTMLResponse)
def debug_page(request: Request, book_id: int | None = None):
    repo, _ = runtime()
    books = repo.list_books("active")
    selected = book_id or (books[0].id if books else None)
    return templates.TemplateResponse(request, "debug.html", {"request": request, "books": books, "book": repo.get_book(selected) if selected else None, "controls": repo.get_book_controls(selected) if selected else None, "result": None, "sample": ""})


@app.post("/debug", response_class=HTMLResponse)
def run_debug(request: Request, book_id: int = Form(...), sample: str = Form("")):
    repo, _ = runtime()
    books = repo.list_books("active")
    controls = repo.get_book_controls(book_id)
    guard_result = TextGuard().check_body(sample)
    char_count = len("".join(sample.split()))
    target_min = int(controls["settings"].get("target_chars_min", 2200))
    target_max = int(controls["settings"].get("target_chars_max", 3200))
    problems = list(guard_result.problems)
    if char_count < target_min:
        problems.append(f"字数不达标：当前约 {char_count} 字，低于最低目标 {target_min} 字。")
    if char_count > target_max + 600:
        problems.append(f"字数失控：当前约 {char_count} 字，高于最高目标 {target_max} 字太多。")
    editorial = EditorialDepartment(repo, TextGuard())
    if controls["settings"].get("pov_policy", "third_limited") == "third_limited" and editorial._looks_first_person(sample):
        problems.append("人称风险：当前设置为第三人称有限视角，正文疑似第一人称旁白。")
    story_problem = editorial._story_shape_problem(sample)
    if story_problem:
        problems.append(story_problem)
    result = {"passed": not problems, "problems": problems, "char_count": char_count}
    return templates.TemplateResponse(request, "debug.html", {"request": request, "books": books, "book": repo.get_book(book_id), "controls": controls, "result": result, "sample": sample})


@app.get("/exports", response_class=HTMLResponse)
def exports_page(request: Request):
    repo, _ = runtime()
    return templates.TemplateResponse(request, "exports.html", {"request": request, "exports": repo.list_exports(), "books": repo.list_books("active")})


@app.post("/books/{book_id}/exports/docx")
def export_docx(request: Request, book_id: int):
    repo, pipe = runtime()
    bodies = repo.list_chapter_bodies(book_id, status="human_confirmed")
    if not bodies:
        return render_error(request, "没有人工确认正文，不能导出 DOCX，也不能写回连续性工作室。")
    book = repo.get_book(book_id)
    path = _writable_output_path(f"book_{book_id}.docx")
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(book.title if book else "未命名作品", 0)
        if book and book.premise:
            doc.add_paragraph(book.premise)
        for row in bodies:
            doc.add_heading(row["title"], level=1)
            for para in row["body"].splitlines():
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(path)
    except ImportError:
        _write_minimal_docx(path, book.title if book else "未命名作品", book.premise if book else "", bodies)
    record = repo.create_export_record(book_id, "docx", str(path))
    export_id = int(record["id"])
    repo.mark_exported(book_id, export_id)
    pipe.continuity.writeback_from_export(book_id, export_id)
    return RedirectResponse("/exports", status_code=303)


@app.get("/exports/{export_id}/download")
def download_export(export_id: int):
    repo, _ = runtime()
    record = next((item for item in repo.list_exports() if int(item["id"]) == export_id), None)
    if record is None:
        return {"error": "not found"}
    path = Path(record["file_path"])
    return FileResponse(path, filename=path.name)


@app.get("/departments/{department}", response_class=HTMLResponse)
def department_page(request: Request, department: str, book_id: int | None = None):
    repo, _ = runtime()
    books = repo.list_books("active")
    selected = book_id or (books[0].id if books else None)
    return templates.TemplateResponse(request, "department.html", {"request": request, "department": department, "books": books, "book": repo.get_book(selected) if selected else None, "controls": repo.get_book_controls(selected) if selected else None, "plans": repo.list_chapter_plans(selected) if selected else []})


@app.get("/health")
def health():
    repo, _ = runtime()
    return {"status": "ok", "books": len(repo.list_books("active"))}


@app.get("/favicon.ico")
def favicon():
    return FileResponse(PROJECT_ROOT / "app" / "static" / "icons" / "favicon.svg", media_type="image/svg+xml")


def _write_minimal_docx(path: Path, title: str, premise: str, bodies: list[dict[str, object]]) -> None:
    paragraphs = [title, premise]
    for row in bodies:
        paragraphs.append(str(row["title"]))
        paragraphs.extend([part.strip() for part in str(row["body"]).splitlines() if part.strip()])
    document_xml = "".join(f"<w:p><w:r><w:t>{html.escape(text)}</w:t></w:r></w:p>" for text in paragraphs if text)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", """<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>""")
        archive.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>""")
        archive.writestr("word/document.xml", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>{document_xml}<w:sectPr/></w:body></w:document>""")


def _writable_output_path(filename: str) -> Path:
    target = EXPORT_ROOT / filename
    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        probe = target.parent / ".write_probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
        return target
    except OSError:
        fallback = Path(tempfile.gettempdir()) / "fiction_architect_exports"
        fallback.mkdir(parents=True, exist_ok=True)
        return fallback / filename


if __name__ == "__main__":
    import uvicorn

    settings = get_settings()
    uvicorn.run("app.main:app", host=settings.app_host, port=settings.app_port, reload=False)
